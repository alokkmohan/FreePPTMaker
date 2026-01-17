import streamlit as st
import os
import glob
import shutil
from datetime import datetime
from docx import Document
from ppt_to_images import ppt_to_images
from create_ppt import process_script
from ai_ppt_generator import generate_beautiful_ppt
from youtube_script_generator import generate_youtube_script_with_ai
from content_generator import generate_content_from_topic

# Page Configuration
st.set_page_config(
    page_title="AI PowerPoint & YouTube Script Generator",
    page_icon="🎨",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for beautiful UI
st.markdown("""
<style>
    /* Main container */
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
    }
    
    /* Header styling */
    .header-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 20px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.3);
        text-align: center;
    }
    
    .main-title {
        font-size: 3rem;
        font-weight: 800;
        color: white;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }
    
    .sub-title {
        font-size: 1.3rem;
        color: #e0e7ff;
        margin-top: 0.5rem;
        font-weight: 300;
    }
    
    /* Content cards */
    .stApp {
        background: #f0f2f6;
    }
    
    div[data-testid="stVerticalBlock"] > div {
        background: white;
        border-radius: 15px;
        padding: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
    }
    
    /* Buttons */
    .stButton > button {
        border-radius: 10px;
        font-weight: 600;
        border: none;
        padding: 0.75rem 2rem;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    }
    
    /* File uploader */
    .uploadedFile {
        border-radius: 10px;
        background: #f8f9fa;
    }
    
    /* Text areas */
    .stTextArea > div > div > textarea {
        border-radius: 10px;
        border: 2px solid #e0e7ff;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Info boxes */
    .stAlert {
        border-radius: 10px;
        border-left: 4px solid #667eea;
    }
    
    /* Download buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 8px;
    }
    
    /* Mobile responsive */
    @media (max-width: 768px) {
        .main-title {
            font-size: 2rem;
        }
        .sub-title {
            font-size: 1rem;
        }
        .header-container {
            padding: 1rem;
        }
        div[data-testid="stVerticalBlock"] > div {
            padding: 1rem;
        }
    }
    
    /* Feature cards */
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 4px solid #667eea;
    }
</style>
""", unsafe_allow_html=True)

# Beautiful Header
st.markdown("""
<div class="header-container">
    <h1 class="main-title">🎨 AI PowerPoint Generator</h1>
    <p class="sub-title">Transform Ideas into Beautiful Presentations & YouTube Scripts</p>
</div>
""", unsafe_allow_html=True)

# Create directories if they don't exist
os.makedirs("input", exist_ok=True)
os.makedirs("output/slides", exist_ok=True)

# Input Method Selection
st.subheader("📝 Step 1: Input Your Script")
input_method = st.radio(
    "Choose input method:",
    ["Upload File", "Paste Text"],
    horizontal=True
)

script_content = None

if input_method == "Upload File":
    uploaded_file = st.file_uploader(
        "Upload your script file (TXT, DOC, DOCX, PDF)",
        type=["txt", "doc", "docx", "pdf", "md"]
    )
    if uploaded_file:
        # Read file content based on type
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == "txt" or file_type == "md":
            script_content = uploaded_file.read().decode('utf-8')
        elif file_type == "docx":
            # For DOCX, we'll save and read using python-docx
            temp_path = os.path.join("input", "temp.docx")
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.read())
            doc = Document(temp_path)
            script_content = "\n".join([para.text for para in doc.paragraphs])
        elif file_type == "pdf":
            st.warning("PDF support coming soon! Please convert to TXT first.")
        else:
            script_content = uploaded_file.read().decode('utf-8')
        
        if script_content:
            st.success("✅ File uploaded successfully!")
            with st.expander("Preview uploaded content"):
                st.text_area("Content Preview", script_content, height=200, disabled=True)

else:  # Paste Text
    st.markdown("**Input Type:**")
    input_type = st.radio(
        "What are you providing?",
        ["Complete Article/Script", "Topic Only (AI will generate content)"],
        horizontal=True,
        help="Choose 'Topic Only' if you want AI to write the content for you"
    )
    
    if input_type == "Topic Only (AI will generate content)":
        topic = st.text_input(
            "Enter your topic:",
            placeholder="Example: AI use in government offices"
        )
        
        if topic:
            st.info(f"📝 Topic: {topic}")
            st.caption("AI will generate detailed content on this topic and create PPT")
            script_content = f"TOPIC:{topic}"  # Special marker for topic mode
    else:
        script_content = st.text_area(
            "Paste your complete article/script here:",
            height=300,
            placeholder="यहाँ अपनी पूरी script paste करें..."
        )
    
    if script_content:
        st.success(f"✅ Input received")

# Process Button
if script_content:
    st.markdown("---")
    st.subheader("🚀 Step 2: Configure Output")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        generate_ppt = st.checkbox("Generate PowerPoint", value=True)
    with col2:
        use_ai = st.checkbox("Use AI Enhancement", value=False, help="Use Groq AI for smart content structuring")
    with col3:
        generate_youtube_script = st.checkbox("Generate YouTube Script", value=True)
    
    # AI customization prompt
    ai_instructions = ""
    if use_ai:
        st.markdown("**🤖 AI Instructions (Optional):**")
        ai_instructions = st.text_area(
            "Tell AI how to structure your presentation:",
            placeholder="Example: 'Make it engaging for students' or 'Focus on technical details' or 'Keep it simple and concise'",
            height=80,
            help="AI will use these instructions to better structure your content"
        )
        if ai_instructions:
            st.success(f"✅ AI will follow your instructions")
    
    # Design options
    if generate_ppt:
        st.markdown("**🎨 Design Theme:**")
        col_a, col_b, col_c, col_d = st.columns(4)
        with col_a:
            theme_ocean = st.button("🌊 Ocean", use_container_width=True)
        with col_b:
            theme_forest = st.button("🌲 Forest", use_container_width=True)
        with col_c:
            theme_sunset = st.button("🌅 Sunset", use_container_width=True)
        with col_d:
            theme_corp = st.button("💼 Corporate", use_container_width=True)
        
        # Set theme in session state
        if theme_ocean:
            st.session_state['theme'] = 'ocean'
        elif theme_forest:
            st.session_state['theme'] = 'forest'
        elif theme_sunset:
            st.session_state['theme'] = 'sunset'
        elif theme_corp:
            st.session_state['theme'] = 'corporate'
        
        selected_theme = st.session_state.get('theme', 'corporate')
        st.info(f"Selected theme: **{selected_theme.title()}**")
    
    if st.button("🎬 Generate Beautiful Files", type="primary", use_container_width=True):
        # Store AI instructions in session state
        st.session_state['ai_instructions'] = ai_instructions
        
        # Check if it's topic mode (AI content generation)
        if script_content.startswith("TOPIC:"):
            topic = script_content.replace("TOPIC:", "").strip()
            
            with st.spinner(f"🤖 AI is generating detailed content on: {topic}..."):
                try:
                    # Generate full article from topic
                    generated_content = generate_content_from_topic(topic, ai_instructions)
                    script_content = generated_content
                    st.success("✅ Content generated successfully!")
                    
                    # Show preview
                    with st.expander("📄 View Generated Content"):
                        st.text_area("Generated Article", generated_content, height=300)
                except Exception as e:
                    st.error(f"❌ Content generation failed: {str(e)}")
                    st.stop()
        
        # Create unique output folder with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = os.path.join("output", f"output_{timestamp}")
        os.makedirs(output_folder, exist_ok=True)
        
        # Set timestamp in session state
        st.session_state['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        st.session_state['output_folder'] = output_folder
        
        # Save script to file
        script_path = os.path.join("input", "RawScript.txt")
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(script_content)
        
        # Also save script to output folder
        script_backup = os.path.join(output_folder, "script.txt")
        shutil.copy(script_path, script_backup)
        
        st.markdown("---")
        st.subheader("📦 Results")
        st.info(f"📁 Output saved in: {output_folder}")
        
        # Generate PPT
        if generate_ppt:
            selected_theme = st.session_state.get('theme', 'corporate')
            user_instructions = st.session_state.get('ai_instructions', '')
            
            with st.spinner(f"Creating beautiful {selected_theme} themed presentation..."):
                ppt_path = os.path.join(output_folder, "presentation.pptx")
                try:
                    # Use AI-powered generator
                    success = generate_beautiful_ppt(
                        script_content, 
                        ppt_path, 
                        color_scheme=selected_theme,
                        use_ai=use_ai,
                        ai_instructions=user_instructions
                    )
                except Exception as e:
                    st.error(f"❌ Error during PPT generation: {str(e)}")
                    success = False
                
                if success:
                    st.success("✅ PowerPoint generated successfully!")
                    
                    # Download PPT
                    with open(ppt_path, "rb") as f:
                        st.download_button(
                            label="⬇️ Download PowerPoint (PPTX)",
                            data=f,
                            file_name="presentation.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            use_container_width=True
                        )
                    
                    # Generate Images from PPT
                    with st.spinner("Converting slides to images..."):
                        img_success = ppt_to_images(ppt_path, output_dir=output_folder)
                        images = sorted([f for f in os.listdir(output_folder) if f.startswith("slide_") and f.endswith(".png")])
                        
                        if images:
                            st.success(f"✅ {len(images)} slide images generated!")
                            
                            # Display images in grid
                            st.markdown("#### Preview Slides")
                            cols = st.columns(3)
                            for idx, img in enumerate(images):
                                with cols[idx % 3]:
                                    img_path = os.path.join(output_folder, img)
                                    st.image(img_path, caption=f"Slide {idx+1}", use_container_width=True)
                                    with open(img_path, "rb") as f:
                                        st.download_button(
                                            label=f"⬇️ {img}",
                                            data=f,
                                            file_name=img,
                                            mime="image/png",
                                            key=f"img_{idx}",
                                            use_container_width=True
                                        )
                        else:
                            st.info("Images not generated. LibreOffice may not be installed.")
                else:
                    st.error("❌ PPT generation failed. Check your script format.")
        
        # Generate YouTube Script DOC
        if generate_youtube_script:
            with st.spinner("Creating professional YouTube script..."):
                # Generate AI-enhanced YouTube script
                youtube_script = generate_youtube_script_with_ai(script_content)
                
                doc_path = os.path.join(output_folder, "youtube_script.docx")
                doc = Document()
                
                # Add title with styling
                title = doc.add_heading("YouTube Script", 0)
                title.alignment = 1  # Center alignment
                
                # Add metadata
                metadata = doc.add_paragraph(f"Generated on: {st.session_state.get('timestamp', 'N/A')}")
                metadata.alignment = 1
                doc.add_paragraph()
                
                # Add horizontal line
                doc.add_paragraph("_" * 80)
                doc.add_paragraph()
                
                # Process and add YouTube script with proper formatting
                lines = youtube_script.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Handle different markdown headers
                    if line.startswith('# '):
                        doc.add_heading(line.replace('# ', ''), 0)
                    elif line.startswith('## '):
                        heading = doc.add_heading(line.replace('## ', ''), 1)
                        heading.alignment = 0
                    elif line.startswith('### '):
                        doc.add_heading(line.replace('### ', ''), 2)
                    elif line.startswith('**') and line.endswith('**'):
                        # Bold text
                        para = doc.add_paragraph()
                        run = para.add_run(line.strip('*'))
                        run.bold = True
                    elif line.startswith('- ') or line.startswith('• '):
                        # Bullet point
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line == '---':
                        # Horizontal line
                        doc.add_paragraph('_' * 80)
                    else:
                        # Normal paragraph
                        if line:
                            doc.add_paragraph(line)
                
                # Save document
                doc.save(doc_path)
                st.success("✅ Professional YouTube script created!")
                
                # Show preview
                with st.expander("📺 Preview YouTube Script"):
                    st.markdown(youtube_script)
                
                # Download button
                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="⬇️ Download YouTube Script (DOC)",
                        data=f,
                        file_name="youtube_script.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
    <p>💡 Tip: Upload any text file or paste your script directly. Get PPT presentations and formatted YouTube scripts!</p>
</div>
""", unsafe_allow_html=True)
