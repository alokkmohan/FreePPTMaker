import streamlit as st
import os
import glob
import shutil
from datetime import datetime
import json
import base64
from docx import Document
import time
try:
    from ppt_to_images import ppt_to_images
    PPT_TO_IMAGES_AVAILABLE = True
except:
    PPT_TO_IMAGES_AVAILABLE = False
try:
    import pdfplumber
    PDF_SUPPORT_AVAILABLE = True
except:
    PDF_SUPPORT_AVAILABLE = False
from create_ppt import process_script
from ai_ppt_generator import generate_beautiful_ppt
from youtube_script_generator import generate_youtube_script_with_ai
from content_generator import generate_content_from_topic
from image_generator import get_slide_image, download_image, search_image
# Claude integration
from claude_ppt_generator import create_ppt_from_file, create_ppt_from_topic
# === CLAUDE FILE UPLOAD SECTION ===
def add_claude_file_upload_section():
    st.markdown("### 🤖 Upload Documents for Professional PPT Generation")
    st.markdown("Upload one or more Word, PDF, or text files to create a professional presentation.")
    uploaded_files = st.file_uploader(
        "Upload your document(s)",
        type=['docx', 'pdf', 'txt', 'md'],
        help="Upload Word, PDF, or text files for professional PPT generation",
        accept_multiple_files=True
    )
    if uploaded_files:
        temp_dir = "temp_uploads"
        os.makedirs(temp_dir, exist_ok=True)
        file_paths = []
        for uploaded_file in uploaded_files:
            file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            file_paths.append(file_path)
        st.success(f"✅ {len(file_paths)} file(s) uploaded: {', '.join([os.path.basename(fp) for fp in file_paths])}")
        col1, col2 = st.columns(2)
        with col1:
            ppt_style = st.selectbox(
                "Presentation Style",
                ["professional", "government", "corporate", "technical"],
                help="Choose the style of your presentation"
            )
            min_slides = st.slider("Minimum Slides", 5, 15, 10)
        with col2:
            audience = st.selectbox(
                "Target Audience",
                ["general", "executives", "technical", "government"],
                help="Who will view this presentation?"
            )
            max_slides = st.slider("Maximum Slides", 10, 25, 15)
        custom_instructions = st.text_area(
            "Additional Instructions (Optional)",
            placeholder="E.g., Focus on data analysis, Include case studies, etc.",
            height=100
        )
        presenter_name = st.text_input(
            "Presenter Name (Optional)",
            placeholder="Your name"
        )
        if st.button("🎨 Generate Professional PPT", type="primary", use_container_width=True):
            with st.spinner("🤖 Analyzing your documents and creating a professional presentation..."):
                output_folder = "outputs"
                os.makedirs(output_folder, exist_ok=True)
                output_path = os.path.join(output_folder, f"presentation_{int(time.time())}.pptx")
                try:
                    # Pass the list of file_paths to your PPT generation logic
                    success = create_ppt_from_file(
                        file_path=file_paths,
                        output_path=output_path,
                        style=ppt_style,
                        min_slides=min_slides,
                        max_slides=max_slides,
                        audience=audience,
                        presenter=presenter_name,
                        custom_instructions=custom_instructions
                    )
                    if success:
                        st.success("✅ Professional PowerPoint created successfully!")
                        with open(output_path, "rb") as f:
                            ppt_data = f.read()
                        st.download_button(
                            label="📥 Download PowerPoint",
                            data=ppt_data,
                            file_name="professional_presentation.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            use_container_width=True,
                            type="primary"
                        )
                        st.info("💡 Your presentation has been generated with professional formatting!")
                    else:
                        st.error("❌ Failed to generate presentation. Please try again.")
                except Exception as e:
                    st.error(f"❌ Error generating presentation:\n\n{str(e)}")
                    st.info("📌 Debug info has been logged. Please check your document and try again.")

# Simple visitor counter and conversion tracking
def update_visitor_count():
    """Track and update visitor count"""
    counter_file = "visitor_count.json"
    
    # Initialize session-based tracking
    if 'visitor_counted' not in st.session_state:
        try:
            # Read current count
            if os.path.exists(counter_file):
                with open(counter_file, 'r') as f:
                    data = json.load(f)
                    count = data.get('total_visits', 0)
                    ppt_count = data.get('ppt_generated', 0)
            else:
                count = 0
                ppt_count = 0
            
            # Increment count
            count += 1
            
            # Save updated count
            with open(counter_file, 'w') as f:
                json.dump({
                    'total_visits': count, 
                    'ppt_generated': ppt_count,
                    'last_visit': datetime.now().isoformat()
                }, f)
            
            st.session_state['visitor_counted'] = True
            st.session_state['visit_count'] = count
        except:
            st.session_state['visit_count'] = 0
    
    return st.session_state.get('visit_count', 0)

def track_ppt_generation():
    """Track successful PPT generation for conversion rate"""
    counter_file = "visitor_count.json"
    
    # Only count once per session
    if 'ppt_tracked' not in st.session_state:
        try:
            # Read current data
            if os.path.exists(counter_file):
                with open(counter_file, 'r') as f:
                    data = json.load(f)
            else:
                data = {'total_visits': 0, 'ppt_generated': 0}
            
            # Increment PPT count
            data['ppt_generated'] = data.get('ppt_generated', 0) + 1
            data['last_ppt'] = datetime.now().isoformat()
            
            # Save updated data
            with open(counter_file, 'w') as f:
                json.dump(data, f)
            
            st.session_state['ppt_tracked'] = True
        except:
            pass

def get_conversion_stats():
    """Get conversion rate statistics"""
    counter_file = "visitor_count.json"
    try:
        if os.path.exists(counter_file):
            with open(counter_file, 'r') as f:
                data = json.load(f)
                visits = data.get('total_visits', 0)
                ppts = data.get('ppt_generated', 0)
                if visits > 0:
                    conversion_rate = (ppts / visits) * 100
                else:
                    conversion_rate = 0
                return visits, ppts, conversion_rate
    except:
        pass
    return 0, 0, 0

# Page Configuration
st.set_page_config(
    page_title="TEXT to PPT Generator - AI Powered",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for beautiful, mobile-responsive UI
st.markdown("""
<style>
    /* CSS Variables for theming */
    :root {
        --bg-primary: linear-gradient(to bottom, #f5f7fa 0%, #e8ecf1 100%);
        --bg-card: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%);
        --text-primary: #2d3748;
        --text-secondary: #718096;
        --border-color: #e8ecf1;
    }
    
    [data-theme="dark"] {
        --bg-primary: linear-gradient(to bottom, #1a202c 0%, #2d3748 100%);
        --bg-card: linear-gradient(145deg, #2d3748 0%, #374151 100%);
        --text-primary: #f7fafc;
        --text-secondary: #cbd5e0;
        --border-color: #4a5568;
    }
    
    /* Main background */
    .stApp {
        background: var(--bg-primary);
    }
    
    /* Header styling */
    .header-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem 1.5rem;
        border-radius: 25px;
        margin-bottom: 2rem;
        box-shadow: 0 15px 35px rgba(102, 126, 234, 0.3);
        text-align: center;
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    .main-title {
        font-size: 2rem;
        font-weight: 900;
        color: white;
        margin: 0 0 0.3rem 0;
        text-shadow: 2px 2px 8px rgba(0,0,0,0.2);
        letter-spacing: -1px;
    }
    
    .sub-title {
        font-size: 1rem;
        color: #f0f4ff;
        margin: 0;
        font-weight: 400;
        opacity: 0.95;
    }
    
    /* Feature cards */
    .feature-card {
        background: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%);
        padding: 2rem;
        border-radius: 18px;
        box-shadow: 0 8px 20px rgba(0,0,0,0.08);
        margin: 1rem 0;
        border: 2px solid #e8ecf1;
        transition: all 0.3s ease;
        text-align: center;
    }
    
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 30px rgba(102, 126, 234, 0.2);
        border-color: #667eea;
    }
    
    .feature-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
    }
    
    .feature-title {
        font-size: 1.4rem;
        font-weight: 700;
        color: #2d3748;
        margin: 0.5rem 0;
    }
    
    .feature-desc {
        color: #718096;
        font-size: 1rem;
        margin: 0;
    }
    
    /* Buttons */
    .stButton > button {
        border-radius: 15px;
        font-weight: 600;
        border: 2px solid #e8ecf1;
        padding: 2rem 1.5rem;
        transition: all 0.3s ease;
        font-size: 1.1rem;
        background: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%);
        color: #2d3748;
        box-shadow: 0 8px 20px rgba(0,0,0,0.08);
        white-space: pre-line;
        height: auto;
        line-height: 1.6;
    }
    
    .stButton > button:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 30px rgba(102, 126, 234, 0.2);
        border-color: #667eea;
    }
    
    /* Primary button */
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.8rem 2.5rem;
        white-space: normal;
        line-height: normal;
    }
    
    /* Menu buttons - consistent size */
    button[data-testid="baseButton-secondary"] {
        padding: 2rem 1.5rem !important;
        font-size: 1.1rem !important;
        background: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%) !important;
        border: 2px solid #e8ecf1 !important;
        color: #2d3748 !important;
        transition: all 0.3s ease !important;
    }
    
    button[data-testid="baseButton-secondary"]:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 10px 25px rgba(102, 126, 234, 0.15) !important;
        border-color: #667eea !important;
    }
    
    /* Active menu button states */
    .menu-active-ppt button[data-testid="baseButton-secondary"]:first-child {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        border: 3px solid #667eea !important;
        box-shadow: 0 12px 30px rgba(102, 126, 234, 0.4) !important;
        transform: scale(1.02) !important;
        font-weight: 700 !important;
    }
    
    .menu-active-youtube button[data-testid="baseButton-secondary"]:last-child {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%) !important;
        color: white !important;
        border: 3px solid #f5576c !important;
        box-shadow: 0 12px 30px rgba(245, 87, 108, 0.4) !important;
        transform: scale(1.02) !important;
        font-weight: 700 !important;
    }
    
    /* File uploader */
    .uploadedFile {
        border-radius: 12px;
        background: #f8f9fa;
        padding: 1rem;
    }
    
    /* Text areas */
    .stTextArea > div > div > textarea {
        border-radius: 12px;
        border: 2px solid #e2e8f0;
        font-size: 1rem;
        padding: 1rem;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Text inputs */
    .stTextInput > div > div > input {
        border-radius: 10px;
        border: 2px solid #e2e8f0;
        padding: 0.8rem;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Info/Success boxes */
    .stAlert {
        border-radius: 12px;
        border-left: 5px solid #667eea;
    }
    
    /* Download buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        font-weight: 600;
    }
    
    /* Section headers - works in both light and dark mode */
    h2, h3 {
        color: var(--text-primary) !important;
        font-weight: 700 !important;
    }
    
    h4, h5, h6 {
        color: var(--text-primary) !important;
        font-weight: 600 !important;
    }
    
    /* Step indicators */
    .step-indicator {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 0.5rem 1.2rem;
        border-radius: 20px;
        font-weight: 700;
        display: inline-block;
        margin-bottom: 1rem;
        box-shadow: 0 4px 10px rgba(102, 126, 234, 0.3);
    }
    
    /* Theme buttons */
    .theme-button {
        padding: 1rem;
        border-radius: 12px;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s;
    }
    
    /* Mobile responsive */
    @media (max-width: 768px) {
        .main-title {
            font-size: 1.6rem;
        }
        .sub-title {
            font-size: 0.9rem;
        }
        .header-container {
            padding: 1.2rem 1.2rem;
        }
        .feature-card {
            padding: 1.5rem;
        }
        .feature-icon {
            font-size: 2.5rem;
        }
        .stButton > button {
            padding: 0.7rem 1.5rem;
            font-size: 1rem;
        }
    }
    
    @media (max-width: 480px) {
        .main-title {
            font-size: 1.8rem;
        }
        .sub-title {
            font-size: 0.95rem;
        }
        .header-container {
            padding: 1.5rem 1rem;
            border-radius: 15px;
        }
        .feature-card {
            padding: 1rem;
        }
    }
    
    /* Footer */
    .footer {
        text-align: center;
        padding: 1.5rem 1.5rem;
        color: white;
        font-size: 0.9rem;
        margin-top: 3rem;
        margin-left: 0;
        margin-right: 0;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 25px;
        box-shadow: 0 15px 35px rgba(102, 126, 234, 0.3);
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    .update-info {
        background: linear-gradient(145deg, #ffffff 0%, #f8f9ff 100%);
        border: 2px solid #667eea;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 2rem auto;
        max-width: 100%;
        box-shadow: 0 8px 20px rgba(102, 126, 234, 0.15);
    }
    
    .update-info h3 {
        color: #667eea;
        font-size: 1.2rem;
        margin: 0 0 0.5rem 0;
        font-weight: 700;
    }
    
    .update-info p {
        margin: 0.5rem 0;
        color: #4a5568;
        font-size: 1rem;
    }
    
    .update-info .timestamp {
        color: #764ba2;
        font-weight: 700;
        font-size: 1.1rem;
        font-family: 'Courier New', monospace;
    }
    
    /* Reset button styling */
    .stButton > button[data-testid="baseButton-secondary"].reset-btn {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%) !important;
        color: white !important;
        border: 2px solid #f5576c !important;
        font-weight: 600 !important;
        padding: 0.6rem 1.2rem !important;
        border-radius: 12px !important;
        box-shadow: 0 6px 15px rgba(245, 87, 108, 0.25) !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button[data-testid="baseButton-secondary"].reset-btn:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 20px rgba(245, 87, 108, 0.35) !important;
    }
</style>
""", unsafe_allow_html=True)

# Minimal Sidebar - only reset button
with st.sidebar:
    if st.button("🔄 Reset App", key="reset_btn", use_container_width=True, help="Clear all data and start fresh"):
        # Clear all session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

st.markdown("""
<style>
    /* Hide default sidebar header */
    section[data-testid="stSidebar"] > div:first-child {
        padding-top: 2rem;
    }
    
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #f8f9ff 0%, #ffffff 100%);
    }
    
    section[data-testid="stSidebar"] h3 {
        color: #667eea;
        font-weight: 700;
    }
</style>
""", unsafe_allow_html=True)

# Beautiful Header
st.markdown("""
<div class="header-container">
    <div class="main-title">📊 TEXT to PPT Generator</div>
    <div class="sub-title">Transform Your Text into Beautiful PowerPoint Presentations</div>
</div>
""", unsafe_allow_html=True)

# Create directories if they don't exist
os.makedirs("input", exist_ok=True)
os.makedirs("output/slides", exist_ok=True)

# Main Menu Selection
# Initialize session state for main menu - auto-select PowerPoint
if 'main_menu' not in st.session_state:
    st.session_state['main_menu'] = 'ppt'  # Auto-select PowerPoint

# Initialize current_tab if not exists
if 'current_tab' not in st.session_state:
    st.session_state['current_tab'] = 'powerpoint'

if 'input_method' not in st.session_state:
    st.session_state['input_method'] = None

# Three Main Menu Buttons (Horizontal)
st.markdown("""
<style>
.stButton > button {
    width: 100% !important;
    min-height: 60px !important;
    font-size: 1.1rem !important;
    font-weight: 600 !important;
    border-radius: 12px !important;
    background: linear-gradient(90deg, #667eea 0%, #a7bfe8 100%) !important;
    color: #fff !important;
    border: 2px solid #e0e0e0 !important;
    transition: all 0.3s ease !important;
    padding: 15px !important;
    margin-bottom: 0.5rem !important;
}
.stButton > button:nth-child(2) {
    background: linear-gradient(90deg, #f7971e 0%, #ffd200 100%) !important;
    color: #fff !important;
}
.stButton > button:nth-child(3) {
    background: linear-gradient(90deg, #f5576c 0%, #f093fb 100%) !important;
    color: #fff !important;
}
.stButton > button:hover {
    border-color: #764ba2 !important;
    opacity: 0.92;
}
</style>
""", unsafe_allow_html=True)

menu_cols = st.columns(3)

with menu_cols[0]:
    if st.button("📊 Choose Input Method for PowerPoint", key="menu_powerpoint", use_container_width=True):
        st.session_state['current_tab'] = 'powerpoint'
        st.rerun()
with menu_cols[1]:
    if st.button("❓ Help & Instructions", key="menu_help", use_container_width=True):
        st.session_state['current_tab'] = 'help'
        st.rerun()
with menu_cols[2]:
    if st.button("☕ Support This Work", key="menu_donate", use_container_width=True):
        st.session_state['current_tab'] = 'donate'
        st.rerun()

# Show content based on selected tab
if st.session_state['current_tab'] == 'help':
    # Help content
    st.markdown("## ℹ️ Help & Instructions")
    st.markdown("""
**How to use this tool:**

1. **Choose Input Method:** Select how you want to provide content (upload, paste, or topic).
2. **Configure Presentation:** Set slide count, theme, and other options.
3. **Generate PPT:** Click the generate button and download your presentation.

**Features:**
- Claude AI-powered professional PPT generation
- Supports Word, PDF, and text files
- Hindi & English content supported
- Custom themes and slide counts

For more details, see the README or contact the developer.
    """)

elif st.session_state['current_tab'] == 'donate':
    # Donate content with smaller centered QR
    col_left, col_center, col_right = st.columns([1.5, 1, 1.5])
    with col_center:
        st.markdown("<h2 style='text-align: center;'>💝 Support This Work</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #764ba2; font-size: 16px;'>If you like this tool and find it useful, consider supporting</p>", unsafe_allow_html=True)
        qr_path = os.path.join(os.getcwd(), "QR-Code.jpeg")
        if os.path.exists(qr_path):
            st.image(qr_path, width=180, caption="Scan to Donate")
        else:
            st.info("QR Code not found. Please add QR-Code.jpeg to project folder.")
        st.markdown("<p style='text-align: center; font-size: 14px;'>UPI • PhonePe • Google Pay • Paytm</p>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center;'><b>Thank you for your support! ☕</b></p>", unsafe_allow_html=True)

# PowerPoint Tab content (default)
if st.session_state['current_tab'] == 'powerpoint':
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📁 Upload Multiple\nDocuments\n\nUpload TXT, DOCX, MD, PDF files", 
                     key="upload_method_btn_1", use_container_width=True, help="Upload one or multiple documents"):
            st.session_state['input_method'] = 'upload'
    with col2:
        if st.button("✍️ Paste Content\n\nCopy and paste your\nready-made content", 
                     key="paste_method_btn_1", use_container_width=True, help="Paste your prepared content"):
            st.session_state['input_method'] = 'paste_article'
    with col3:
        if st.button("🎯 Write Topic\n\nLet AI generate content\nfrom your topic", 
                     key="topic_method_btn_1", use_container_width=True, help="AI will generate presentation"):
            st.session_state['input_method'] = 'write_topic'
    st.markdown("")


# Initialize script_content outside conditionals
script_content = None


# Only show upload section if on PowerPoint tab
if st.session_state.get('current_tab') == 'powerpoint' and st.session_state.get('input_method') == 'upload':
    add_claude_file_upload_section()

elif st.session_state.get('input_method') == 'write_topic':
    st.markdown("### 🎯 Enter Your Topic")
    st.markdown("AI will generate a detailed article and create presentation from your topic")
    topic_input = st.text_area(
        label="Enter your topic",
        placeholder="Example:\n• Artificial Intelligence in Healthcare\n• Climate Change and Its Effects\n• Future of Electric Vehicles\n• Digital Marketing Strategies 2025"
    )
    use_professional = st.checkbox("🤖 Use Professional PPT Generation", value=False, help="Enable to use advanced AI-powered generation for this topic")
    if topic_input:
        if st.session_state.get('last_topic') != topic_input:
            st.session_state['confirmed_content'] = f"TOPIC:{topic_input}"
            st.session_state['last_topic'] = topic_input
        st.success(f"✅ Topic confirmed: **{topic_input}**")
        st.info("📊 Your AI-powered presentation is being prepared... scroll down to configure your presentation!")
    else:
        st.warning("⬇️ Enter a topic above to get started")
    script_content = st.session_state.get('confirmed_content') if st.session_state.get('confirmed_content', '').startswith('TOPIC:') else None
    # Professional topic generation UI
    if use_professional and topic_input:
        st.markdown("---")
        st.markdown("### 🤖 Generate Professional PPT from Topic")
        ai_choice = st.selectbox(
            "Select AI Model",
            ["Deepseek", "Gemini", "Groq", "Hugging Face"],
            key="ai_model_choice"
        )
        col1, col2 = st.columns(2)
        with col1:
            ppt_style = st.selectbox(
                "Presentation Style",
                ["professional", "government", "corporate", "technical"],
                key="topic_style"
            )
            min_slides = st.slider("Minimum Slides", 5, 15, 10, key="topic_min")
        with col2:
            audience = st.selectbox(
                "Target Audience",
                ["general", "executives", "technical", "government"],
                key="topic_aud"
            )
            max_slides = st.slider("Maximum Slides", 10, 25, 15, key="topic_max")
        custom_instructions = st.text_area(
            "Additional Instructions (Optional)",
            placeholder="E.g., Include recent statistics, Focus on implementation, etc.",
            height=100,
            key="topic_instr"
        )
        if st.button(f"🚀 Generate with {ai_choice}", type="primary", key="topic_btn"):
            with st.spinner(f"🤖 {ai_choice} is researching and creating your presentation..."):
                output_folder = "outputs"
                os.makedirs(output_folder, exist_ok=True)
                import re
                safe_title = re.sub(r'[^-\w\s-]', '', topic_input)[:50]
                safe_title = re.sub(r'[-\s]+', '_', safe_title)
                output_path = os.path.join(output_folder, f"{safe_title}_{ai_choice.lower()}.pptx")
                # TODO: Call the appropriate backend function for Deepseek, Gemini, Groq, or Hugging Face here
                st.info(f"(Backend call for {ai_choice} will be implemented here.)")

elif st.session_state.get('input_method') == 'paste_article':
    st.markdown("### 📝 Paste Your Content")
    
    # Auto-save draft feature
    if 'draft_content' not in st.session_state:
        st.session_state['draft_content'] = ""
    
    article_input = st.text_area(
        "",
        value=st.session_state.get('draft_content', ''),
        height=300,
        placeholder="Paste your content here...\nYou can paste in multiple parts - the text will accumulate.\nSupports multiple languages including Hindi, English, etc.",
        help="Paste your ready-to-convert content. You can paste multiple times before submitting.",
        label_visibility="collapsed",
        key="article_text",
        on_change=lambda: st.session_state.update({'draft_content': st.session_state.article_text})
    )
    
    # Auto-save indicator
    if article_input:
        st.caption("💾 Auto-saved draft")
        
        # Content Analytics for pasted content
        word_count = len(article_input.split())
        char_count = len(article_input)
        estimated_slides = max(5, min(20, word_count // 100))
        
        # Quality Indicators
        avg_word_length = sum(len(word) for word in article_input.split()) / max(word_count, 1)
        readability_score = "Good" if 4 <= avg_word_length <= 7 else "Complex" if avg_word_length > 7 else "Simple"
        
        col_stat1, col_stat2, col_stat3 = st.columns(3)
        with col_stat1:
            st.metric("📝 Words", f"{word_count:,}")
        with col_stat2:
            st.metric("📊 Est. Slides", estimated_slides)
        with col_stat3:
            st.metric("🎯 Readability", readability_score)
    
    # AI Enhancement option
    st.markdown("####")
    enhance_with_ai = st.checkbox(
        "🤖 AI Enhancement - Improve and structure content with AI",
        value=True,
        help="Enable this to let AI enhance, restructure and improve your content before creating presentation"
    )
    
    if enhance_with_ai:
        st.info("💡 AI will enhance and restructure your content for better presentation")
    else:
        st.info("📄 Will create presentation directly from your pasted content")
    
    col_confirm, col_clear = st.columns([3, 1])
    with col_confirm:
        if article_input and st.button("✅ Confirm Content and Proceed", type="primary", use_container_width=True):
            # Store both content and enhancement preference
            if enhance_with_ai:
                st.session_state['confirmed_content'] = f"ENHANCE:{article_input}"
            else:
                st.session_state['confirmed_content'] = article_input
            st.session_state['ai_enhancement'] = enhance_with_ai
            st.success(f"✅ Content confirmed ({len(article_input)} characters) - {'AI Enhancement ON' if enhance_with_ai else 'Direct Conversion'}")
    
    with col_clear:
        if st.button("🗑️ Clear", use_container_width=True):
            st.session_state['draft_content'] = ""
            st.rerun()
    
    script_content = st.session_state.get('confirmed_content')
    # Filter out TOPIC: prefix for this check
    if script_content and script_content.startswith('TOPIC:'):
        script_content = None

else:
    # Initialize script_content if no input method selected
    script_content = None

# Process if content is available
if script_content and st.session_state.get('main_menu'):
    st.markdown("---")
    
    # Set generate flags - PowerPoint only
    generate_ppt = True
    generate_youtube_script = False
    
    # AI Enhancement setting based on user choice
    ai_enhancement_enabled = st.session_state.get('ai_enhancement', False)
    use_ai = ai_enhancement_enabled  # Only use AI if user opted for enhancement
    ai_instructions = ""
    
    # Smart slide count calculation
    word_count = len(script_content.split())
    if word_count <= 300:
        auto_slides = "5-8 slides"
        default_min, default_max = 5, 8
    elif word_count <= 800:
        auto_slides = "10-15 slides"
        default_min, default_max = 10, 15
    else:
        auto_slides = "15-25 slides"
        default_min, default_max = 15, 25
    
    # Slide count selection
    st.markdown("### 📊 Presentation Length")
    
    slide_preference = st.radio(
        "Choose presentation style:",
        ["✨ Auto (Recommended)", "🎯 Quick (5-8 slides)", "📄 Standard (10-15 slides)", "📚 Detailed (15-25 slides)"],
        horizontal=True,
        help="Auto mode intelligently adjusts based on content length"
    )

    if slide_preference == "✨ Auto (Recommended)":
        min_slides, max_slides = default_min, default_max
        st.session_state['slide_preference'] = 'auto'
    elif slide_preference == "🎯 Quick (5-8 slides)":
        min_slides, max_slides = 5, 8
        st.session_state['slide_preference'] = 'quick'
    elif slide_preference == "📄 Standard (10-15 slides)":
        min_slides, max_slides = 10, 15
        st.session_state['slide_preference'] = 'standard'
    else:  # Detailed
        min_slides, max_slides = 15, 25
        st.session_state['slide_preference'] = 'detailed'

    # Store in session state
    st.session_state['min_slides'] = min_slides
    st.session_state['max_slides'] = max_slides
    
    # Theme selection (only for PPT mode)
    if generate_ppt:
        st.markdown("####")
        st.markdown("### 🎨 Choose Design Theme")
        
        col_a, col_b, col_c, col_d, col_e = st.columns(5)
        
        with col_a:
            if st.button("🌊 Ocean", use_container_width=True, help="Professional blue theme"):
                st.session_state['theme'] = 'ocean'
                st.session_state['template_path'] = None
        with col_b:
            if st.button("🌲 Forest", use_container_width=True, help="Natural green theme"):
                st.session_state['theme'] = 'forest'
                st.session_state['template_path'] = None
        with col_c:
            if st.button("🌅 Sunset", use_container_width=True, help="Warm orange theme"):
                st.session_state['theme'] = 'sunset'
                st.session_state['template_path'] = None
        with col_d:
            if st.button("💼 Corporate", use_container_width=True, help="Classic business theme"):
                st.session_state['theme'] = 'corporate'
                st.session_state['template_path'] = None
        with col_e:
            if st.button("🎯 EG Theme", use_container_width=True, help="EG custom red theme"):
                st.session_state['theme'] = 'eg'
                st.session_state['template_path'] = None
        
        # Display selected options
        selected_theme = st.session_state.get('theme', 'corporate')
        selected_pref = st.session_state.get('slide_preference', 'auto')
        min_s = st.session_state.get('min_slides', 10)
        max_s = st.session_state.get('max_slides', 15)
        
        pref_names = {
            'auto': '✨ Auto (Recommended)',
            'quick': '🎯 Quick',
            'standard': '📄 Standard',
            'detailed': '📚 Detailed'
        }
        
        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.info(f"🎨 Theme: **{selected_theme.upper()}**")
        with col_info2:
            st.info(f"📊 Style: **{pref_names.get(selected_pref, 'Standard')}** ({min_s}-{max_s} slides)")
    
    st.markdown("####")
    if st.button(f"🚀 Generate PowerPoint", type="primary", use_container_width=True):
        st.session_state['ai_instructions'] = ai_instructions

        # Store original topic for PPT title
        original_topic = None

        # Dynamic word/section requirements based on slide preference
        min_s = st.session_state.get('min_slides', 10)
        max_s = st.session_state.get('max_slides', 15)
        slide_pref = st.session_state.get('slide_preference', 'auto')
        if slide_pref == 'quick':
            word_range = "700-1000 words"
            section_range = "5-8 slides/sections"
        elif slide_pref == 'standard':
            word_range = "1200-1800 words"
            section_range = "10-15 slides/sections"
        elif slide_pref == 'detailed':
            word_range = "2500-4000 words"
            section_range = "15-25 slides/sections"
        else:
            word_range = f"{min_s*120}-{max_s*180} words"
            section_range = f"{min_s}-{max_s} slides/sections"

        # Handle topic mode (always uses AI)
        if script_content.startswith("TOPIC:"):
            topic = script_content.replace("TOPIC:", "").strip()
            original_topic = topic  # Store the original topic

            # Progress bar for AI generation
            progress_text = st.empty()
            progress_bar = st.progress(0)

            progress_text.text("🤖 Initializing AI...")
            progress_bar.progress(10)

            with st.spinner(f"🤖 AI is generating detailed content on: **{topic}**..."):
                try:
                    progress_text.text("📝 Generating comprehensive article...")
                    progress_bar.progress(30)

                    generated_content = generate_content_from_topic(topic, ai_instructions, min_s, max_s)

                    progress_bar.progress(80)
                    progress_text.text("✅ Content generated! Structuring slides...")

                    script_content = generated_content
                    progress_bar.progress(100)
                    progress_text.empty()
                    progress_bar.empty()

                    st.success("✅ Content generated successfully!")

                    with st.expander("📄 View Generated Content"):
                        st.text_area("Generated Content", generated_content, height=300, label_visibility="collapsed")
                except Exception as e:
                    progress_bar.empty()
                    progress_text.empty()
                    st.error(f"❌ Content generation failed: {str(e)}")
                    st.stop()

        # Handle AI enhancement mode for pasted content
        elif script_content.startswith("ENHANCE:"):
            original_content = script_content.replace("ENHANCE:", "").strip()

            with st.spinner(f"🤖 AI is enhancing and structuring your content..."):
                try:
                    # Use the content generator to enhance the pasted content
                    enhanced_content = generate_content_from_topic(f"Enhance and restructure this content:\n\n{original_content}", ai_instructions, min_s, max_s)
                    script_content = enhanced_content
                    st.success("✅ Content enhanced successfully!")

                    with st.expander("📄 View Enhanced Content"):
                        st.text_area("Enhanced Content", enhanced_content, height=300, label_visibility="collapsed")
                except Exception as e:
                    st.error(f"❌ Content generation failed: {str(e)}")
                    st.stop()
        
        # Create output folder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = os.path.join("output", f"output_{timestamp}")
        os.makedirs(output_folder, exist_ok=True)
        
        st.session_state['timestamp'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        st.session_state['output_folder'] = output_folder
        
        # Save script
        script_path = os.path.join("input", "RawScript.txt")
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(script_content)
        
        script_backup = os.path.join(output_folder, "script.txt")
        shutil.copy(script_path, script_backup)
        
        st.markdown("---")
        st.markdown('<div class="step-indicator">📦 Results</div>', unsafe_allow_html=True)
        st.markdown("")
        
        # Generate PPT
        if generate_ppt:
            selected_theme = st.session_state.get('theme', 'corporate')
            user_instructions = st.session_state.get('ai_instructions', '')
            
            with st.spinner(f"🎨 Creating beautiful **{selected_theme.upper()}** presentation..."):
                # Generate filename from title
                import re
                if original_topic:
                    safe_title = re.sub(r'[^\w\s-]', '', original_topic)[:50]
                    safe_title = re.sub(r'[-\s]+', '_', safe_title)
                    ppt_filename = f"{safe_title}.pptx"
                else:
                    ppt_filename = "presentation.pptx"
                ppt_path = os.path.join(output_folder, ppt_filename)
                
                # Get template path from session state
                template_path = st.session_state.get('template_path', None)
                
                # Get slide count preferences
                min_slides = st.session_state.get('min_slides', 10)
                max_slides = st.session_state.get('max_slides', 20)
                
                try:
                    success = generate_beautiful_ppt(
                        script_content, 
                        ppt_path, 
                        color_scheme=selected_theme,
                        use_ai=use_ai,
                        ai_instructions=user_instructions,
                        original_topic=original_topic,
                        template_path=template_path,
                        min_slides=min_slides,
                        max_slides=max_slides
                    )
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
                    success = False
                
                if success:
                    # Track conversion
                    track_ppt_generation()
                    
                    st.success("✅ PowerPoint created successfully!")
                    
                    # Store in session state for persistence
                    st.session_state['ppt_generated'] = True
                    st.session_state['ppt_path'] = ppt_path
                    
                    # Add to download history
                    if 'download_history' not in st.session_state:
                        st.session_state['download_history'] = []
                    
                    st.session_state['download_history'].append({
                        'name': 'presentation.pptx',
                        'path': ppt_path,
                        'type': 'PowerPoint',
                        'mime': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                        'time': datetime.now().strftime("%Y-%m-%d %H:%M")
                    })
                    
                    # Download the PPT file
                    with open(ppt_path, "rb") as f:
                        ppt_data = f.read()
                    
                    # Display prominent download button
                    st.markdown("---")
                    col1, col2, col3 = st.columns([1, 2, 1])
                    with col2:
                        st.download_button(
                            label="📥 Download PowerPoint",
                            data=ppt_data,
                            file_name="presentation.pptx",
                            mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                            use_container_width=True,
                            key=f"download_ppt_{int(datetime.now().timestamp())}",
                            type="primary"
                        )
                    st.markdown("---")
                    
                    st.success("✅ PPT Generated Successfully! Click the Download button above.")
                    
                    # Generate images for slide preview
                    st.markdown("---")
                    st.markdown("### 🖼️ Slide Preview")
                    
                    try:
                        if PPT_TO_IMAGES_AVAILABLE:
                            with st.spinner("🖼️ Converting slides to images..."):
                                img_success = ppt_to_images(ppt_path, output_dir=output_folder)
                                images = sorted([f for f in os.listdir(output_folder) if f.startswith("slide_") and f.endswith(".png")])
                                
                                # Store images in session state
                                st.session_state['slide_images'] = images
                                st.session_state['images_folder'] = output_folder
                                
                                if images:
                                    st.success(f"✅ {len(images)} slides generated")
                                    
                                    # Sort images by slide number to ensure correct order
                                    sorted_images = sorted(images, key=lambda x: int(x.split('_')[1].split('.')[0]))
                                    
                                    # Display slides in 3 columns
                                    cols = st.columns([1, 1, 1])
                                    
                                    for idx, img in enumerate(sorted_images):
                                        col_idx = idx % 3
                                        with cols[col_idx]:
                                            img_path = os.path.join(output_folder, img)
                                            if os.path.exists(img_path):
                                                st.markdown(f"**Slide {idx+1}**")
                                                st.image(img_path, use_column_width=True)
                                else:
                                    st.info("📄 Image preview could not be generated")
                        else:
                            st.info("📄 Image conversion not available - PPT is ready to download")
                    except Exception as e:
                        st.info("📄 Image preview could not be generated - PPT is ready to download")
                else:
                    st.error("❌ PPT generation failed")
        
        # Generate YouTube Script
        if generate_youtube_script:
            with st.spinner("🎬 Creating YouTube script..."):
                youtube_script = generate_youtube_script_with_ai(script_content)
                
                doc_path = os.path.join(output_folder, "youtube_script.docx")
                doc = Document()
                
                title = doc.add_heading("YouTube Script", 0)
                title.alignment = 1
                
                metadata = doc.add_paragraph(f"Generated: {st.session_state.get('timestamp', 'N/A')}")
                metadata.alignment = 1
                doc.add_paragraph()
                doc.add_paragraph("_" * 80)
                doc.add_paragraph()
                
                # Format content
                lines = youtube_script.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.startswith('# '):
                        doc.add_heading(line.replace('# ', ''), 0)
                    elif line.startswith('## '):
                        doc.add_heading(line.replace('## ', ''), 1)
                    elif line.startswith('### '):
                        doc.add_heading(line.replace('### ', ''), 2)
                    elif line.startswith('**') and line.endswith('**'):
                        para = doc.add_paragraph()
                        run = para.add_run(line.strip('*'))
                        run.bold = True
                    elif line.startswith('- ') or line.startswith('• '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line == '---':
                        doc.add_paragraph('_' * 80)
                    else:
                        if line:
                            doc.add_paragraph(line)
                
                doc.save(doc_path)
                st.success("✅ YouTube script created!")
                
                # Store in session state
                st.session_state['youtube_generated'] = True
                st.session_state['youtube_path'] = doc_path
                st.session_state['youtube_script'] = youtube_script
                
                # Add to download history
                if 'download_history' not in st.session_state:
                    st.session_state['download_history'] = []
                
                st.session_state['download_history'].append({
                    'name': 'youtube_script.docx',
                    'path': doc_path,
                    'type': 'YouTube Script',
                    'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'time': datetime.now().strftime("%Y-%m-%d %H:%M")
                })
                
                with st.expander("📺 Preview YouTube Script"):
                    st.markdown(youtube_script)
                
                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="📥 Download YouTube Script (DOCX)",
                        data=f,
                        file_name="youtube_script.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

# Footer - appears on all pages
st.markdown("""
<div class='footer'>
    All systems operational | Hindi & English support | 5-25 slides generation<br>
    Developed by Alok Mohan. verson v2
</div>
""", unsafe_allow_html=True)
